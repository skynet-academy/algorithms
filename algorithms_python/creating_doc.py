from docx import Document
from docx.text.run import Font, Run
from docx.shared import Inches, RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE


class PsychotestDocument:

    def __init__(self, patient, visit, age, birth):
        self.document = Document()
        self.patient = patient
        self.visit = visit
        self.patient_age = age
        self.patient_birth = birth
        self.doc_template()
        self.patient_data()
        self.document.save("my_doc.docx")
 
    def doc_template(self):
        self.title = self.document.add_paragraph("Психологическое обследование по научной теме «Ранняя диагностика БА»")
        self.title.style = self.document.styles.add_style("Style Name", WD_STYLE_TYPE.PARAGRAPH)
        self.font = self.title.style.font
        self.font.name = "Times New Roman"
        self.font.size = Pt(12)
        self.font.bold = True
        #font.italic = True
        # #font.underline = True
        self.font.color.rgb = RGBColor(0, 0, 0)    # this is black


    #def patient_data(self, document, patient, visit, age, birth):
    def patient_data(self):
        self.surname = self.document.add_paragraph("ФИО пациента:")
        self.surname.add_run(self.patient)
        self.patient_data = self.document.add_paragraph("Дата обследования: ")
        self.patient_data.add_run(self.visit)
        self.patient_birthday = self.document.add_paragraph("Дата рождения: ")
        self.patient_birthday.add_run(self.patient_birth)
        self.patient_age = self.document.add_paragraph("Возраст пациента: ")
        #self.patient_age.add_run(self.patient_age)
    
first = PsychotestDocument('nicolas', '10/10/2020', '32', '10/10/2020')







